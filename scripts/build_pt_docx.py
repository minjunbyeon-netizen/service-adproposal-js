# -*- coding: utf-8 -*-
"""영산대 PT 발표키트 Word 문서 생성 — 시나리오 + 강조포인트 두 docx 빌드.

군더더기(호흡 곡선·만트라·비상 시나리오·신체 동선)는 제외.
시간배분·침묵 시간·핵심 멘트·Q&A만 유지.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(r"C:/Users/USER/Desktop/영산대_PT발표키트_V101")
FONT = "맑은 고딕"


def set_font(run, name=FONT, size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, level=1, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    size = size or sizes.get(level, 11)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, size=10, bold=False, color=None, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# helper
# ============================================================
def _new_doc(margins=(2, 2, 2.2, 2.2)):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(margins[0])
        section.bottom_margin = Cm(margins[1])
        section.left_margin = Cm(margins[2])
        section.right_margin = Cm(margins[3])
    return doc


# ============================================================
# 문서 1: 발표 시나리오 30분
# ============================================================
def build_scenario(doc=None, do_save=True):
    if doc is None:
        doc = _new_doc((2, 2, 2.2, 2.2))

    # 제목
    title = doc.add_paragraph()
    run = title.add_run("영산대 광고대행사 선정 PT — 30분 발표 시나리오")
    set_font(run, size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "2026-05-07(목) 14:00 / 해운대캠퍼스 D동 1층 대회의실 / 30분 발표 + Q&A 10분 / 2개사 중 1번째",
        size=10,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "발표자: 하이브미디어 진성희 차장",
        size=10,
        bold=True,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 시간 배정 총괄
    add_heading(doc, "1. 시간 배정 총괄", level=1)
    add_table(
        doc,
        ["Part", "슬라이드", "시간", "내용"],
        [
            ["Part 1 — 도입·진단", "P1~P11 (11장)", "6:45", "첫 90초 hook 집중"],
            ["Part 2 — 무엇을 기억시킬 것인가", "P12~P36 (25장)", "12:10", "컨셉 + 시안"],
            ["Part 3 — 어떻게 많이 팔 것인가", "P37~P52 (16장)", "7:50", "실행 + 예산"],
            ["Part 4 — 클로징", "P53~P54 (2장)", "2:00", "수미상관 + 인사"],
            ["합계", "54장", "28:45", "30:00까지 1:15 버퍼"],
        ],
        col_widths=[5.5, 3.5, 1.8, 5.5],
    )

    # 2. Part 1
    add_heading(doc, "2. Part 1 — 도입·진단 (P1~P11 / 6분 45초)", level=1)
    add_table(
        doc,
        ["페이지", "시간", "누적", "키 메시지"],
        [
            ["P1 표지 0825503", "0:20", "0:20", "'안녕하십니까. 하이브미디어 진성희입니다. 광고대행사 선정은 두 개로 압축됩니다.'"],
            ["P2 Evaluation Axis", "0:45", "1:05", "'무엇을 기억시킬 것인가. 어떻게 많이 팔 것인가'"],
            ["P3 Market Status", "0:25", "1:30", "'모든 대학은 지금 생존의 시대'"],
            ["P4 Competitive Map", "0:40", "2:10", "'같은 무대에서 각자의 무기로'"],
            ["P5 About Youngsan", "0:25", "2:35", "'영산대학교의 무기는 무엇입니까?' — 침묵 2초"],
            ["P6 Hospitality", "0:20", "2:55", "'호텔관광 특성화 대학이라는 무기가 있습니다'"],
            ["P7 Brand Recall", "0:50", "3:45", "'호텔관광 하면 즉시 영산대학교라고는 못하셨을 겁니다'"],
            ["P8 SERIAL 0825503", "1:20", "5:05", "08 — 25 — 55 — 03 한 숫자씩 1초씩 정지"],
            ["P9 Memory Challenge", "1:10", "6:15", "5초 표시 → 페이드 → 침묵 3초 → '몇 개 기억하십니까?'"],
            ["P10 Core Imprint", "0:30", "6:45", "'숫자를 기억시키는 것. 그게 전부입니다'"],
            ["P11 Company Profile", "0:00*", "6:45", "*P10에서 '저희 하이브미디어가 그 일을 합니다'로 통과"],
        ],
        col_widths=[3.5, 1.5, 1.5, 9.8],
    )

    # 3. Part 2
    add_heading(doc, "3. Part 2 — 무엇을 기억시킬 것인가 (P12~P36 / 12분 10초)", level=1)
    add_table(
        doc,
        ["페이지", "시간", "누적", "키 메시지"],
        [
            ["P12 챕터 표지", "0:10", "6:55", "'1막입니다'"],
            ["P13 To Imprint", "0:45", "7:40", "'기억시켜야 할 것 — 08 연구력 세계 8위 NATIONAL NO.1'"],
            ["P14 Ad Sample", "1:10", "8:50", "'학교명을 가렸습니다. 어디인지 맞추시겠습니까?' — 침묵 2초"],
            ["P15 Memory Test", "0:50", "9:40", "'동신대·수원여대·한국기술교육대 — 모두 1위입니다'"],
            ["P16 Memory Test", "0:20", "10:00", "'기억이 안 되시죠'"],
            ["P17 Why Forgotten", "0:30", "10:30", "'왜 기억되지 않을까요?' — 답하지 말고 다음으로"],
            ["P18 4-Layer Intro", "0:35", "11:05", "'4계층 레이어로 답하겠습니다'"],
            ["P19 Ad Concept", "1:10", "12:15", "컨셉 선언 — 0825503. 진입 3초 + 선언 후 침묵 5초"],
            ["P20 Imprinting 3요소", "0:30", "12:45", "'Brand Image 3가지: 산업·세계·재학'"],
            ["P21 Target First", "0:25", "13:10", "'누구에게 말할지부터 — 3 타겟'"],
            ["P22~P24 Industry 01-03", "1:00", "14:10", "'산업이 먼저 찾는 대학' (각 0:20)"],
            ["P25~P27 Global 01-03", "1:00", "15:10", "'서울 기준으로 고르는 시대는 끝났다' (각 0:20)"],
            ["P28~P30 Student 01-03", "1:00", "16:10", "'재학 중부터 업계인' (각 0:20)"],
            ["P31~P33 Global Recruit", "0:45", "16:55", "베트남·중국·영문 (각 0:15)"],
            ["P34 Transition", "0:20", "17:15", "'인쇄로 각인했다면, 이제 영상입니다'"],
            ["P35 Video Story", "1:00", "18:15", "영상 재생 (없으면 스토리 구술)"],
            ["P36 1 Page Summary", "0:40", "18:55", "'산업이 선택한 대학, 세계가 증명한 숫자'"],
        ],
        col_widths=[3.5, 1.5, 1.5, 9.8],
    )

    # 4. Part 3
    add_heading(doc, "4. Part 3 — 어떻게 많이 팔 것인가 (P37~P52 / 7분 50초)", level=1)
    add_table(
        doc,
        ["페이지", "시간", "누적", "키 메시지"],
        [
            ["P37 챕터 표지", "0:10", "19:05", "'2막입니다'"],
            ["P38 Expansion", "0:30", "19:35", "'3 타겟 → 6 채널로 확장합니다'"],
            ["P39 Matrix", "0:35", "20:10", "표 한 번 짚어주고"],
            ["P40 PRE 채널 전략", "0:35", "20:45", "'PRE — 사전 노출 단계'"],
            ["P41 Influencer Review", "1:00", "21:45", "CPE 320원 — 업계 평균 600~1200원의 절반 이하"],
            ["P42 STEP 1 유튜브", "0:20", "22:05", "'STEP 1 — 자체제작 유튜브'"],
            ["P43 STEP 2 숏폼", "0:20", "22:25", "'STEP 2 — 재학생 숏폼'"],
            ["P44 STEP 3 인쇄", "0:20", "22:45", "'STEP 3 — 인쇄 (입학처 직배포)'"],
            ["P45 STEP 4 디지털 DA", "0:20", "23:05", "'STEP 4 — 디지털 DA'"],
            ["P46 STEP 5 SNS", "0:20", "23:25", "'STEP 5 — SNS'"],
            ["P47 STEP 6 언론", "0:20", "23:45", "'STEP 6 — 언론'"],
            ["P48 Management Plan", "0:30", "24:15", "'전담 인력 3인, 외주 의존 없음'"],
            ["P49 Timeline", "0:25", "24:40", "'2026년 4월 착수, 9월 수시까지 5개월 리드타임'"],
            ["P50 Budget 1.25억", "1:10", "25:50", "'1억 2,500만 원, 이렇게 씁니다' — 항목별 짚기"],
            ["P51 Cost Efficiency", "0:35", "26:25", "'제대로 쓰겠습니다' — CPE 효율 강조"],
            ["P52 Monthly Execution", "0:20", "26:45", "'수시 시즌(9~11월) 집중 편성'"],
        ],
        col_widths=[3.5, 1.5, 1.5, 9.8],
    )

    # 5. Part 4 클로징
    add_heading(doc, "5. Part 4 — 클로징 (P53~P54 / 2분)", level=1)
    add_table(
        doc,
        ["페이지", "시간", "누적", "키 메시지"],
        [
            [
                "P53 Bookend",
                "1:20",
                "28:05",
                "(침묵 5초) → '처음으로 돌아갑니다. 무엇을 기억시킬 것인가. 어떻게 많이 팔 것인가' → '저희의 답: 0825503. 수험생·학부모·유학생'",
            ],
            ["P54 감사합니다", "0:40", "28:45", "'감사합니다' — 인사 후 1초 멈춤"],
        ],
        col_widths=[3.5, 1.5, 1.5, 9.8],
    )

    add_para(doc, "")
    add_para(
        doc,
        "P53 핵심 — 침묵 5초",
        size=11,
        bold=True,
    )
    add_bullet(doc, "P2와 동일한 화면이 뜨는 순간, 평가위원이 '아, P2다'를 인식하는 시간을 준다.")
    add_bullet(doc, "슬라이드가 뜨자마자 말하지 않는다. 이 5초가 PT 전체에서 가장 비싼 순간.")

    # 6. 시간 점검 신호
    add_heading(doc, "6. 시간 점검 신호 (발표자 손목시계)", level=1)
    add_table(
        doc,
        ["누적 시간", "도달 페이지", "늦으면"],
        [
            ["6:45", "P11 끝 (Part 1 완료)", "Part 2 시안(P22~P33) 각 0:05 단축"],
            ["12:45", "P20 끝 (컨셉 선언 후)", "P31~P33 유학생 시안 각 0:10 압축"],
            ["18:55", "P36 끝 (Part 2 완료)", "Part 3 STEP 6개 각 0:15까지 단축"],
            ["25:50", "P50 끝 (예산 설명 완료)", "P51·P52 합쳐 0:30 안에 마무리"],
            ["28:00", "P53 진입", "P53 침묵 5초 → 3초로 단축"],
            ["28:45", "P54 인사 (목표 종료)", "정시 인사"],
        ],
        col_widths=[2.5, 6.5, 7.3],
    )

    add_para(doc, "")
    add_para(
        doc,
        "사수 우선 슬라이드 (시간 절대 단축 금지)",
        size=11,
        bold=True,
    )
    add_bullet(doc, "P8 SERIAL 0825503 — 1:20")
    add_bullet(doc, "P9 Memory Challenge — 1:10")
    add_bullet(doc, "P14 Ad Sample (학교명 가림) — 1:10")
    add_bullet(doc, "P19 컨셉 선언 (침묵 5초) — 1:10")
    add_bullet(doc, "P50 Budget 1.25억 — 1:10")
    add_bullet(doc, "P53 Bookend (침묵 5초) — 1:20")

    if do_save:
        out = OUT_DIR / "02_발표시나리오_30분.docx"
        doc.save(str(out))
        print(f"OK: {out}")
    return doc


# ============================================================
# 문서 2: 강조포인트 + Q&A
# ============================================================
def build_emphasis(doc=None, do_save=True):
    if doc is None:
        doc = _new_doc((2, 2, 2.2, 2.2))

    title = doc.add_paragraph()
    run = title.add_run("영산대 광고대행사 선정 PT — 강조 포인트 + Q&A")
    set_font(run, size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "30분 발표 + Q&A 10분 / 2026-05-07(목) 14:00 / 해운대 D동 / 2개사 중 1번째",
        size=10,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "발표자: 하이브미디어 진성희 차장",
        size=10,
        bold=True,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 강조 7장 한눈 보기
    add_heading(doc, "강조 7장 한눈 보기", level=1)
    add_table(
        doc,
        ["순번", "페이지", "슬라이드", "시간", "배점 매핑"],
        [
            ["1", "P8", "SERIAL 0825503", "1:20", "크리에이티브 30"],
            ["2", "P9", "Memory Challenge", "1:10", "크리에이티브 30"],
            ["3", "P14", "Ad Sample (학교명 가림)", "1:10", "홍보전략 20"],
            ["4", "P19", "Ad Concept 선언", "1:10", "크리에이티브 30"],
            ["5", "P36", "1 Page Summary", "0:40", "크리에이티브 30"],
            ["6", "P50", "Budget 1.25억", "1:10", "매체·예산 20"],
            ["7", "P53", "Bookend (수미상관)", "1:20", "마지막 인상"],
        ],
        col_widths=[1.2, 1.5, 6.5, 1.6, 4.5],
    )

    emphases = [
        {
            "no": "1",
            "page": "P8 — SERIAL No. 0825503 (1분 20초)",
            "why": "PT 전체의 SSOT. 크리에이티브 30점의 가장 직접적 근거. 평가위원이 이 장만 기억해도 다른 업체와 차별화된다.",
            "ment": [
                "(P7에서 자연스럽게 이어) '그래서 영산대학교의 무기를 숫자로 정리해봤습니다.'",
                "(숫자를 일부러 늦게) '공팔 — (1초) — 이공오 — (1초) — 오오 — (1초) — 공삼.'",
                "'08 — 연구력 세계 8위, 국내 사립대 1위.'",
                "'25 — 호텔 총지배인 25명을 배출한 학교.'",
                "'55 — QS 호스피탈리티 분야 세계 55위, 비수도권 1위.'",
                "'03 — 한국 호텔관광 TOP 3.'",
                "'이 숫자들은 영산대학교가 이미 갖고 있는 것들입니다. 광고가 발명하는 게 아닙니다. 꺼내는 겁니다.'",
            ],
            "q": "'세계 8위라는 근거가 무엇입니까?'",
            "a": "THE Impact Rankings 2024 기준입니다. UN SDGs 달성 기여도 평가 지표로, 연구 영향력 항목에서 국내 사립대 1위입니다. 슬라이드 하단에 출처 각주를 명시했고, 모든 광고 소재에도 동일한 출처 표기를 의무화하겠습니다.",
        },
        {
            "no": "2",
            "page": "P9 — Memory Challenge (1분 10초)",
            "why": "크리에이티브 30점 + 발표 초반 장악력. 다른 어느 PT에도 없는 포맷. 청중이 자신의 기억력을 직접 시험하는 순간 우리 컨셉의 정당성이 자동 증명된다.",
            "ment": [
                "(P8 5초 표시 후 페이드아웃)",
                "'방금 보여드린 숫자, 몇 개나 기억하십니까?'",
                "(2초 침묵)",
                "'08, 25, 55, 03. 이 네 숫자를 전부 기억하신 분, 손 들어주시겠습니까?'",
                "(3초 기다림)",
                "'감사합니다. 저희가 오늘 해드릴 일은 딱 하나입니다.'",
                "'영산대학교의 수험생이, 학부모가, 그리고 유학생이 — 이 숫자를 기억하게 만드는 것입니다.'",
            ],
            "q": "'수험생들이 광고에서 숫자를 기억할 수 있을지 의문입니다.'",
            "a": "그래서 광고 단독으로는 안 됩니다. P18에서 설명드릴 4계층 반복 노출 구조가 핵심입니다. 한 번 보여주는 게 아니라 매체마다 같은 숫자를 반복 각인시킵니다. 4번 반복 시 재인률이 단일 노출 대비 6.4배 상승합니다.",
        },
        {
            "no": "3",
            "page": "P14 — Ad Sample 학교명 가림 (1분 10초)",
            "why": "홍보 전략·환경 분석 20점. 차별화 당위성을 가장 직관적으로 증명. 웃음이 나오면 평가위원과 공범 의식이 생긴다.",
            "ment": [
                "'1막 — 무엇을 기억시킬 것인가. 첫 번째 질문입니다.'",
                "'여기 두 대학교 광고가 있습니다. 학교명을 가렸습니다.'",
                "'AAA 대학교, BBB 대학교.'",
                "(2초 침묵) '어느 학교인지 맞추시겠습니까?'",
                "(웃음 유도)",
                "'모르시죠. 저희도 모릅니다. 이 광고들은 실제로 집행된 광고들입니다.'",
                "'지금 시장에 나와 있는 대부분의 대학 광고는 학교명을 바꿔도 구분이 안 됩니다.'",
            ],
            "q": "'경쟁사 광고를 직접 비교하는 게 적절합니까?'",
            "a": "실명을 가린 익명 처리 샘플입니다. 시장 현황 분석의 일부이며, 영산대학교가 차별화해야 하는 이유를 설명하는 데 필요한 근거 자료입니다. 본 PT에서만 사용하고, 실제 광고 집행 시에는 경쟁사 비교 표현을 일절 사용하지 않습니다.",
        },
        {
            "no": "4",
            "page": "P19 — Ad Concept 선언 (1분 10초)",
            "why": "크리에이티브 30점의 정점. 4계층 레이어의 결론이자 모든 시안의 근거. 이 PT의 가장 비싼 1분.",
            "ment": [
                "(P18에서 자연스럽게 이어) '그래서 저희 컨셉은 단 하나입니다.'",
                "(진입 3초 침묵)",
                "'0825503.'",
                "'이 일곱 자리 숫자가 영산대학교의 시리얼 넘버입니다.'",
                "'제품에 시리얼 넘버가 있듯, 이 대학교에는 이 숫자가 있습니다.'",
                "'수험생이 영산대를 검색하기 전에 이 숫자를 먼저 기억하게 만드는 것이 저희의 목표입니다.'",
                "(선언 후 5초 침묵 고정 — 다음 슬라이드로 넘기지 않는다)",
            ],
            "q": "'숫자 광고가 과연 감성적 공감을 이끌어낼 수 있나요? 너무 차갑지 않습니까?'",
            "a": "P22부터 보여드릴 시안에서 확인하실 수 있습니다. Industry, Global, Student 세 타겟별로 각각 다른 감성 카피를 씁니다. 숫자는 뼈대고, 카피가 살입니다. 예를 들어 P23 Industry 02의 카피 '부산을 떠나는 비행기마다, 영산이 타고 있습니다' — 25라는 숫자(총지배인 25명)에서 파생된 감성 표현입니다.",
        },
        {
            "no": "5",
            "page": "P36 — 1 Page Summary (40초)",
            "why": "Part 2 마무리. 평가위원이 크리에이티브 창의성 30점 점수를 머릿속으로 확정하는 타이밍. 한 줄 요약이 평가표 코멘트에 그대로 옮겨진다.",
            "ment": [
                "(시안 12장 직후) '정리하겠습니다.'",
                "'산업이 선택한 대학, 세계가 증명한 숫자.'",
                "'이게 전부입니다.'",
                "'저희가 만들 광고의 모든 카피는 이 한 문장으로 귀결됩니다.'",
                "'1막은 여기서 마칩니다.'",
            ],
            "q": "'이 요약 문장이 다른 대학에도 적용 가능한 일반적 문장 아닙니까?'",
            "a": "'산업이 선택'은 총지배인 25명이라는 수치 근거가 있고, '세계가 증명'은 QS 55위라는 수치 근거가 있습니다. 이 두 수치를 동시에 가진 지방 사립대는 영산대뿐입니다. 일반론처럼 들리지만 각 단어마다 영산대만의 숫자가 박혀 있습니다.",
        },
        {
            "no": "6",
            "page": "P50 — Budget 1.25억 (1분 10초)",
            "why": "매체 전략·예산 적절성 20점 직결. 예산 장이 납득되면 실행력 신뢰가 올라간다.",
            "ment": [
                "'1억 2,500만 원. 이렇게 씁니다.'",
                "'디지털 DA와 SNS가 절반을 넘습니다.'",
                "'인쇄 매체는 의도적으로 줄였습니다.'",
                "'수험생은 유튜브에 있고, 부모는 네이버에 있기 때문입니다.'",
                "'매체를 선택한 게 아니라, 타겟이 있는 곳을 따라간 겁니다.'",
                "'그리고 인플루언서 마케팅에서 CPE 320원 — 업계 평균의 절반 이하입니다.'",
            ],
            "q": "'인쇄 매체 비중이 너무 적은 것 아닙니까?'",
            "a": "대입 수험생 타겟에서 인쇄 접촉률은 디지털의 10% 미만입니다. 다만 P44에 보여드린 고품질 인쇄 시안은 입학처 직배포용으로 별도 편성되어 있습니다. 부울경 고3 담임 약 2,400명에게 직배송하는 구조입니다. 양으로 줄였지만 도달은 정확합니다.",
        },
        {
            "no": "7",
            "page": "P53 — Bookend 수미상관 (1분 20초)",
            "why": "1번째 발표자가 평가위원 뇌리에 끝까지 남는 유일한 길. 배점과 무관하게 마지막 기억을 지배.",
            "ment": [
                "(P52 종료 후 P53 띄움)",
                "(침묵 5초 — 평가위원이 'P2다'를 인식하는 시간)",
                "'처음으로 돌아갑니다.'",
                "'무엇을 기억시킬 것인가. 어떻게 많이 팔 것인가.'",
                "(2초 침묵)",
                "'저희의 답은 변하지 않았습니다.'",
                "'0825503. 이 숫자를 부산 수험생이 먼저 기억하게 만들겠습니다.'",
                "'수험생에게, 학부모에게, 그리고 유학생에게.'",
                "'Social Proof로. Insight Translation으로. Imprinting으로.'",
                "'그래서, 이견없는 영산대학교가 되도록 만들겠습니다.'",
            ],
            "q": "'수미상관 구조가 인상적이긴 한데, 실제 광고 집행 후 측정 지표는 무엇입니까?'",
            "a": "P48 관리 계획에 명시되어 있습니다. 브랜드 서베이 인지도, 원서 접수 전환율, 채널별 CPE — 세 가지 지표를 월별로 보고드립니다. 숫자로 시작했으니 숫자로 결과를 보여드리겠습니다. 분기마다 평가회를 통해 미달성 항목은 다음 분기에 공식 보완 계획서로 제출하겠습니다.",
        },
    ]

    for e in emphases:
        add_heading(doc, f"{e['no']}. {e['page']}", level=2)

        add_para(doc, "왜 중요한가", size=10, bold=True)
        add_para(doc, e["why"], size=10, indent=0.3)

        add_para(doc, "발표 멘트", size=10, bold=True)
        for line in e["ment"]:
            add_bullet(doc, line, size=10, indent=0.5)

        add_para(doc, "예상 질문 → 답변", size=10, bold=True)
        add_para(doc, f"Q. {e['q']}", size=10, indent=0.3, bold=True)
        add_para(doc, f"A. {e['a']}", size=10, indent=0.3)

    # Q&A 10개
    add_heading(doc, "Q&A 예상 질문 10개", level=1)
    add_para(
        doc,
        "답변은 40~50초 / 3~5문장. 길어지면 다음 질문 기회 손실.",
        size=10,
        bold=True,
    )

    qas = [
        {
            "q": "Q1. 1.25억이 적정한 예산 규모입니까? 다른 지방대 대비 어떻습니까?",
            "a": [
                "부울경 지방 사립대 대입 광고 평균 집행액은 1억~2억 구간입니다.",
                "1.25억은 중간값이지만, 디지털 집중 편성으로 CPE 기준 효율은 상위 20%를 목표로 합니다.",
                "인플루언서 CPE 평균 320원은 동종 업계 평균(600~1,200원) 대비 절반 이하입니다.",
                "예산 규모보다 예산 배분 논리가 핵심이며, P50에서 항목별 근거를 확인하셨습니다.",
            ],
        },
        {
            "q": "Q2. 0825503 숫자 각인 컨셉이 실제 광고에서 통한 사례가 있습니까?",
            "a": [
                "국내: 군 복무 기간 '18개월' 캠페인, 보험사 '3330' 상품명 인지율.",
                "해외: BMW M시리즈 숫자 아이덴티티, 나이키 'Air 180'.",
                "교육부 대학광고 수용자 조사 2023: 추상 슬로건 대비 구체 수치의 재인률 2.3배 높음.",
                "영산대는 수치가 이미 있습니다. 만들 필요가 없습니다.",
            ],
        },
        {
            "q": "Q3. 부울경 수험생이 줄어드는 상황에서 지역 타겟 광고가 효과가 있겠습니까?",
            "a": [
                "수험생 감소는 사실이지만, 파이가 줄수록 점유율 싸움이 더 치열해집니다.",
                "3트랙 동시 공략: 부산 재수생 + 경남 고3 + 유학생 유입 (P38).",
                "P31~P33 유학생 시안은 수도권 경쟁 없는 블루오션입니다.",
                "감소하는 파이 안에서 점유율을 높이는 것 + 새 파이(유학생)를 여는 것을 병행합니다.",
            ],
        },
        {
            "q": "Q4. 인플루언서 마케팅의 신뢰도 문제(허위·과장)를 어떻게 해결합니까?",
            "a": [
                "협찬 공개 의무화 이후 팔로워 수가 아닌 실질 참여율(CPE)이 지표입니다.",
                "저희 CPE 320원은 순수 참여 기반 — 팔로워 수 기준 선발 아닙니다.",
                "영산대 호텔관광 졸업생 현직자 크리에이터 우선 발굴 — 광고가 아니라 경험담.",
                "모든 영상에 협찬 표기 의무, 허위 발견 시 즉시 송출 중단 + 페널티 조항.",
            ],
        },
        {
            "q": "Q5. 서울 수도권 대학과의 경쟁에서 어떤 우위가 있습니까?",
            "a": [
                "경쟁 구도 자체를 재설정합니다 — 수도권 vs 지방이 아니라 '호텔관광 카테고리 1위' 싸움.",
                "세종대·경희대 호텔관광학과 대비 총지배인 25명은 국내 최다 수준입니다.",
                "부산이라는 지리적 약점을 '현장 인접성'으로 뒤집은 게 P25~P27 Global 시안.",
                "서울 가지 않아도 세계 호텔 체인에 취업할 수 있다는 메시지입니다.",
            ],
        },
        {
            "q": "Q6. 전담 인력 구성이 어떻게 됩니까? 이 프로젝트에 몇 명이 붙습니까?",
            "a": [
                "P48 관리 계획대로 크리에이티브 디렉터 1, AE 1, 디지털 퍼포먼스 1 전담 체계입니다.",
                "외주 의존 없이 인하우스 제작 역량으로 소재 교체 사이클 4주 → 2주 단축.",
                "월별 보고 + 긴급 소재 수정 48시간 이내 납기 보장.",
                "비상 시 본사 크리에이터 풀 8명까지 즉시 가용합니다.",
            ],
        },
        {
            "q": "Q7. QS 세계 55위 수치가 어느 분야 기준인지 정확히 설명해주십시오.",
            "a": [
                "QS World University Rankings by Subject — Hospitality & Leisure Management 분야 기준입니다.",
                "영산대 공식 발표 자료 기반이며, 광고 사용 시 출처와 기준 병기 의무화.",
                "모든 슬라이드와 인쇄물에 'QS by Subject 2024 기준' 명시 → 소비자 오인 방지.",
                "수치의 진위보다, 어떻게 정확하게 커뮤니케이션할지가 저희 역할입니다.",
            ],
        },
        {
            "q": "Q8. 경쟁사 대비 이 회사만의 차별점이 무엇입니까?",
            "a": [
                "오늘 보여드린 체험형 Memory Test 포맷은 저희 고유 방법론입니다.",
                "시안 제작이 템플릿이 아니라 데이터 기반 수치 포지셔닝 설계에서 출발합니다.",
                "대입 광고 분야 자체 인플루언서 CPE 벤치마크 데이터 보유.",
                "이 캠페인이 왜 영산대에만 맞는가를 숫자로 설명할 수 있는 팀입니다.",
            ],
        },
        {
            "q": "Q9. 2026년 4월 착수면 입시 시즌까지 시간이 촉박하지 않습니까?",
            "a": [
                "P49 타임라인 — 4~5월 소재 제작, 6~8월 디지털 집중, 9~11월 수시 시즌 집중.",
                "수시 1차 9월 → 5개월 리드타임 확보.",
                "계약 후 2주 이내 크리에이티브 브리프 확정, 4주 이내 첫 시안 납기 가능.",
                "늦은 게 아니라, 지금이 정확한 타이밍입니다.",
            ],
        },
        {
            "q": "Q10. PT에서 제시한 시안과 실제 집행이 같습니까?",
            "a": [
                "오늘 보여드린 시안은 전부 실제 집행 가능한 소재 기준으로 제작 — 무드보드 아닙니다.",
                "예산 항목도 매체사 단가 역산 — 희망 예산이 아니라 실측 단가입니다.",
                "계약 시 월별 KPI 목표값을 수치로 합의, 매월 대시보드로 공유.",
                "PT에서 말한 것과 집행이 일치하는지를 저희는 숫자로 증명합니다.",
            ],
        },
    ]

    for qa in qas:
        add_para(doc, qa["q"], size=11, bold=True)
        for line in qa["a"]:
            add_bullet(doc, line, size=10, indent=0.5)

    # Q&A 답변 시 금기
    add_heading(doc, "Q&A 답변 시 금기 사항", level=1)
    add_table(
        doc,
        ["금기", "이유", "대안"],
        [
            ["'글쎄요…' / '잘 모르겠는데요…'", "신뢰 즉시 붕괴", "'확인 후 24시간 이내 서면 답변 드리겠습니다'"],
            ["'그건 다음 단계에서…'", "회피로 인식", "'지금 답변드리고, 디테일은 계약 후 추가 협의'"],
            ["마케팅 전문용어", "Consideration Set / Reverse Recall 등", "일상 한국어로 풀어쓰기"],
            ["클리셰", "학령인구 감소 / 글로벌 경쟁력 / MZ세대", "구체 수치로 대체"],
            ["답변 1분 초과", "다른 질문 기회 손실", "3~5문장에서 의도적으로 자른다"],
        ],
        col_widths=[5.0, 5.5, 6.5],
    )

    if do_save:
        out = OUT_DIR / "03_강조포인트_QnA.docx"
        doc.save(str(out))
        print(f"OK: {out}")
    return doc


# ============================================================
# 문서 3: 단상 컨닝 페이퍼 (A4 1장)
# ============================================================
def build_cheatsheet(doc=None, do_save=True):
    if doc is None:
        doc = _new_doc((1.2, 1.2, 1.5, 1.5))

    title = doc.add_paragraph()
    run = title.add_run("영산대 PT — 단상 컨닝 페이퍼 (목표 종료 28:45)")
    set_font(run, size=16, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "5/7(목) 14:00 / 해운대 D동 / 1번째 발표 / 30분 + Q&A 10분 / 발표자 진성희 차장", size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Part 시간
    add_heading(doc, "Part 시간 (손목시계 기준)", level=2, size=12)
    add_table(
        doc,
        ["체크 시각", "도달 페이지", "남은 시간"],
        [
            ["6:45", "P11 끝 (Part 1 완료)", "23:15"],
            ["12:45", "P20 끝 (컨셉 선언 후)", "17:15"],
            ["18:55", "P36 끝 (Part 2 완료)", "11:05"],
            ["25:50", "P50 끝 (예산 완료)", "4:55"],
            ["28:00", "P53 진입 (클로징)", "2:00"],
            ["28:45", "P54 인사 (목표 종료)", "0:00"],
        ],
        col_widths=[3.5, 8.5, 3.5],
    )

    # 사수 6장
    add_heading(doc, "사수 6장 — 시간 절대 단축 금지", level=2, size=12)
    add_table(
        doc,
        ["페이지", "시간", "한 줄 핵심"],
        [
            ["P8 SERIAL 0825503", "1:20", "08—25—55—03 한 숫자씩 1초 정지"],
            ["P9 Memory Challenge", "1:10", "'몇 개 기억하십니까?' → 손 들기 → 3초 침묵"],
            ["P14 Ad Sample", "1:10", "'학교명 가렸습니다. 맞추시겠습니까?' → 침묵 → 웃음"],
            ["P19 컨셉 선언", "1:10", "'0825503' 선언 후 침묵 5초 (이 5초가 가장 비싸다)"],
            ["P50 Budget 1.25억", "1:10", "'타겟이 있는 곳을 따라간 겁니다' + CPE 320원"],
            ["P53 Bookend", "1:20", "(침묵 5초) → '처음으로 돌아갑니다'"],
        ],
        col_widths=[5.5, 1.5, 8.5],
    )

    # Q&A 키워드 매칭
    add_heading(doc, "Q&A 키워드 매칭 (질문 들으면 즉시 매칭)", level=2, size=12)
    add_table(
        doc,
        ["키워드", "답변 첫 문장"],
        [
            ["예산 적정성", "Q1 → '부울경 평균 1~2억 중 중간값. CPE 효율 상위 20% 목표'"],
            ["숫자 광고 사례", "Q2 → '국내 18개월/3330, 해외 BMW M/Air 180. 재인률 2.3배'"],
            ["수험생 감소", "Q3 → '파이 줄수록 점유율 싸움. 3트랙 + 유학생 블루오션'"],
            ["인플루언서 신뢰", "Q4 → 'CPE 기반. 졸업생 현직자 우선. 협찬 표기 의무'"],
            ["수도권 경쟁", "Q5 → '경쟁 구도 재설정. 호텔관광 카테고리 1위 싸움'"],
            ["전담 인력", "Q6 → 'CD 1·AE 1·퍼포먼스 1. 인하우스. 사이클 4주→2주'"],
            ["QS 55위 근거", "Q7 → 'Subject 2024 Hospitality 분야. 출처 표기 의무화'"],
            ["차별점", "Q8 → 'Memory Test 고유 포맷. 데이터 기반 수치 포지셔닝'"],
            ["일정 촉박", "Q9 → '5개월 리드타임. 4주 내 첫 시안 납기'"],
            ["PT vs 실제", "Q10 → '실측 단가 역산. 월별 KPI 수치 합의 + 대시보드'"],
            ["모르는 질문", "→ '확인 후 24시간 이내 서면 답변 드리겠습니다'"],
        ],
        col_widths=[3.5, 12.0],
    )

    if do_save:
        out = OUT_DIR / "05_컨닝페이퍼_A4.docx"
        doc.save(str(out))
        print(f"OK: {out}")
    return doc


# ============================================================
# 문서 4: 첫 90초 + 마지막 90초 암기 카드
# ============================================================
def build_memorize(doc=None, do_save=True, out_name="04_암기카드_첫90초_마지막90초.docx"):
    if doc is None:
        doc = _new_doc((2, 2, 2.2, 2.2))

    title = doc.add_paragraph()
    run = title.add_run("암기 카드 — 첫 90초 + 마지막 90초")
    set_font(run, size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "이 두 부분은 절대 종이 보지 않고 외운다. 발표 당일 들고 가지 않음 (리허설 전용).",
        size=10,
        bold=True,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 첫 90초
    add_heading(doc, "첫 90초 (P1~P9, 발표 시작 ~ 6:15)", level=1)

    add_para(doc, "P1 — 표지 (0:00~0:20)", size=11, bold=True)
    add_para(doc, "'안녕하십니까. 하이브미디어 진성희입니다. 광고대행사 선정은 두 개로 압축됩니다.'", size=11, indent=0.5)

    add_para(doc, "P2 — Evaluation Axis (0:20~1:05)", size=11, bold=True)
    add_para(doc, "'무엇을 기억시킬 것인가. 어떻게 많이 팔 것인가. 나머지는 수단입니다.'", size=11, indent=0.5)

    add_para(doc, "P3~P7 — 진단 (1:05~3:45)", size=11, bold=True)
    add_para(doc, "'모든 대학은 지금 생존의 시대입니다. 같은 무대에서 각자의 무기로 싸우고 있습니다.'", size=11, indent=0.5)
    add_para(doc, "'그래서, 영산대학교의 무기는 무엇입니까? — 호텔관광 특성화 대학입니다.'", size=11, indent=0.5)
    add_para(doc, "'호텔관광 하면 떠오르시나요? 즉시 영산대학교라고는 못하셨을 겁니다.'", size=11, indent=0.5)

    add_para(doc, "P8 — SERIAL 0825503 (3:45~5:05) ★ 가장 느리게", size=11, bold=True)
    add_para(doc, "'그래서 영산대학교의 무기를 숫자로 정리해봤습니다.'", size=11, indent=0.5)
    add_para(doc, "'공팔 — (1초) — 이공오 — (1초) — 오오 — (1초) — 공삼.'", size=11, indent=0.5)
    add_para(doc, "'08 연구력 세계 8위. 25 호텔 총지배인 25명. 55 QS 세계 55위. 03 한국 TOP 3.'", size=11, indent=0.5)
    add_para(doc, "'이 숫자들은 영산대학교가 이미 갖고 있는 것들입니다. 광고가 발명하는 게 아닙니다. 꺼내는 겁니다.'", size=11, indent=0.5)

    add_para(doc, "P9 — Memory Challenge (5:05~6:15) ★ 체험 hook", size=11, bold=True)
    add_para(doc, "'방금 보여드린 숫자, 몇 개나 기억하십니까?'", size=11, indent=0.5)
    add_para(doc, "(2초 침묵)", size=10, indent=0.5, color=RGBColor(0x88, 0x88, 0x88))
    add_para(doc, "'08, 25, 55, 03. 이 네 숫자를 전부 기억하신 분, 손 들어주시겠습니까?'", size=11, indent=0.5)
    add_para(doc, "(3초 기다림 — 손 적게 올라와도 OK)", size=10, indent=0.5, color=RGBColor(0x88, 0x88, 0x88))
    add_para(doc, "'감사합니다. 저희가 오늘 해드릴 일은 딱 하나입니다.'", size=11, indent=0.5)
    add_para(doc, "'영산대학교의 수험생이, 학부모가, 그리고 유학생이 — 이 숫자를 기억하게 만드는 것입니다.'", size=11, indent=0.5)

    # 마지막 90초
    add_heading(doc, "마지막 90초 (P53~P54, 28:00 ~ 28:45)", level=1)

    add_para(doc, "P53 — Bookend 진입 (28:00~28:05) ★ 침묵 5초 사수", size=11, bold=True)
    add_para(doc, "(P52 종료 후 P53 띄우자마자 5초 침묵 — 평가위원이 'P2다'를 인식하는 시간)", size=10, indent=0.5, color=RGBColor(0x88, 0x88, 0x88))

    add_para(doc, "P53 — 본 클로징 (28:05~28:45)", size=11, bold=True)
    add_para(doc, "'처음으로 돌아갑니다.'", size=11, indent=0.5)
    add_para(doc, "'무엇을 기억시킬 것인가. 어떻게 많이 팔 것인가.'", size=11, indent=0.5)
    add_para(doc, "(2초 침묵)", size=10, indent=0.5, color=RGBColor(0x88, 0x88, 0x88))
    add_para(doc, "'저희의 답은 변하지 않았습니다.'", size=11, indent=0.5)
    add_para(doc, "'0825503. 이 숫자를 부산 수험생이 먼저 기억하게 만들겠습니다.'", size=11, indent=0.5)
    add_para(doc, "'수험생에게, 학부모에게, 그리고 유학생에게.'", size=11, indent=0.5)
    add_para(doc, "'Social Proof로. Insight Translation으로. Imprinting으로.'", size=11, indent=0.5)
    add_para(doc, "'그래서, 이견없는 영산대학교가 되도록 만들겠습니다.'", size=11, indent=0.5)

    add_para(doc, "P54 — 인사 (28:45)", size=11, bold=True)
    add_para(doc, "'감사합니다.'", size=11, indent=0.5, bold=True)
    add_para(doc, "(인사 후 1초 멈춤 → 평가위원 한 명씩 시선 → 단상 정리)", size=10, indent=0.5, color=RGBColor(0x88, 0x88, 0x88))

    # 리허설 가이드
    add_heading(doc, "리허설 방법", level=1)
    add_bullet(doc, "혼자 큰 소리로 첫 90초 5회, 마지막 90초 5회 반복.")
    add_bullet(doc, "스마트폰 녹음 → 실제 시간 측정. 각각 90초 ±5초 안에 들어와야 함.")
    add_bullet(doc, "거울 앞에서 시선 처리 연습 — 슬라이드와 청중 사이 50:50.")
    add_bullet(doc, "P9 '손 들어주시겠습니까?' 후 침묵 3초 — 시계 보지 말고 마음속으로 1001~1003.")
    add_bullet(doc, "P19 '0825503' 선언 후 침묵 5초 — 1001~1005 천천히.")
    add_bullet(doc, "P53 진입 침묵 5초 — 같은 방법.")

    if do_save:
        out = OUT_DIR / out_name
        doc.save(str(out))
        print(f"OK: {out}")
    return doc


# ============================================================
# 문서 5: 위기 대응 1장
# ============================================================
def build_crisis(doc=None, do_save=True):
    if doc is None:
        doc = _new_doc((1.5, 1.5, 2, 2))

    title = doc.add_paragraph()
    run = title.add_run("위기 대응 — 발표 중 5가지 시나리오")
    set_font(run, size=16, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "단상에 펴놓되 보지 않을 자료. 그러나 1번이라도 읽어두면 당황하지 않는다.", size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. 시간 5분 지연 (25분 시점에 P40 미만)", level=2)
    add_bullet(doc, "P42~P47 STEP 6개 통째로 0:30 안에 채널명만 호명: '유튜브—숏폼—인쇄—DA—SNS—언론'.")
    add_bullet(doc, "P48~P49 합쳐 0:30. 'P50으로 바로 가겠습니다'.")
    add_bullet(doc, "P50·P53만 시간 사수. 나머지 압축 OK.")

    add_heading(doc, "2. 영상 재생 오류 (P35)", level=2)
    add_bullet(doc, "당황하지 말고 30초 구술로 대체: '한 학생이 부산 영산대를 졸업하고, 세계 호텔 체인의 총지배인이 됩니다. 그 길이 이미 25명에게 열려 있습니다.'")
    add_bullet(doc, "다음 슬라이드 P36으로 자연스럽게 넘어가며 '영상이 재생되지 않는군요. 다음 장으로 가겠습니다'.")
    add_bullet(doc, "사후에 '영상 본 자료는 별도 송부 드리겠습니다' 구두 약속.")

    add_heading(doc, "3. P9에서 평가위원이 손 들지 않음", level=2)
    add_bullet(doc, "절대 당황하지 말 것. 손 안 올라오는 게 더 좋다 (메시지 강화).")
    add_bullet(doc, "'그렇죠. 그게 광고의 현실입니다. 그래서 저희가 이 자리에 있습니다.'")
    add_bullet(doc, "곧바로 P10으로 넘어가며 침묵 없이 진행.")

    add_heading(doc, "4. 시간초과 임박 (29:30 도달)", level=2)
    add_bullet(doc, "P53 침묵 5초 → 즉시 0초로 단축. 슬라이드 뜨면 바로 '처음으로 돌아갑니다'.")
    add_bullet(doc, "P53 본문도 압축: '0825503. 수험생·학부모·유학생. 이견없는 영산대학교. 감사합니다.'")
    add_bullet(doc, "29:55에는 무조건 '감사합니다' 발화. 정시 인사가 시간초과보다 낫다.")

    add_heading(doc, "5. Q&A 모르는 질문", level=2)
    add_bullet(doc, "절대 '글쎄요' / '잘 모르겠는데요' 금지.")
    add_bullet(doc, "정형 멘트: '좋은 질문 감사드립니다. 정확한 데이터 확인 후 24시간 이내 서면으로 답변 드리겠습니다.'")
    add_bullet(doc, "그 후 즉시 다음 질문 받기. 답변 시도 절대 금지 (틀리면 더 큰 손해).")

    add_para(doc, "")
    add_para(
        doc,
        "마지막 원칙: 어떤 위기든 표정·자세 변화 없이 진행. 평가위원은 '발표자가 위기를 어떻게 다루는가'도 평가한다.",
        size=10,
        bold=True,
    )

    if do_save:
        out = OUT_DIR / "07_위기대응_1장.docx"
        doc.save(str(out))
        print(f"OK: {out}")
    return doc


# ============================================================
# 통합 문서 A: 정독용 (시나리오 + 강조포인트 + Q&A)
# ============================================================
def build_combined_reading():
    doc = _new_doc((2, 2, 2.2, 2.2))
    build_scenario(doc=doc, do_save=False)
    doc.add_page_break()
    build_emphasis(doc=doc, do_save=False)
    out = OUT_DIR / "02_정독용_시나리오_강조_QnA.docx"
    doc.save(str(out))
    print(f"OK: {out}")


# ============================================================
# 통합 문서 B: 단상용 (컨닝페이퍼 + 위기대응)
# ============================================================
def build_combined_dais():
    doc = _new_doc((1.2, 1.2, 1.5, 1.5))
    build_cheatsheet(doc=doc, do_save=False)
    doc.add_page_break()
    build_crisis(doc=doc, do_save=False)
    out = OUT_DIR / "03_단상용_컨닝페이퍼_위기대응.docx"
    doc.save(str(out))
    print(f"OK: {out}")


def cleanup_old():
    """이전 8파일 구조의 잔존 파일 제거."""
    olds = [
        "02_발표시나리오_30분.docx",
        "03_강조포인트_QnA.docx",
        "04_PT_URL.txt",
        "05_컨닝페이퍼_A4.docx",
        "06_암기카드_첫90초_마지막90초.docx",
        "07_위기대응_1장.docx",
    ]
    for old in olds:
        p = OUT_DIR / old
        if p.exists():
            p.unlink()
            print(f"REMOVED: {old}")


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    cleanup_old()
    build_combined_reading()
    build_combined_dais()
    build_memorize()  # 04_암기카드_첫90초_마지막90초.docx
    print("\nDONE.")
    for f in sorted(OUT_DIR.iterdir()):
        print(f"  {f.name}")

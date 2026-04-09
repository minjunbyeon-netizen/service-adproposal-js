import logging
import pdfplumber

logger = logging.getLogger(__name__)


def extract_text_from_pdf(filepath):
    """PDF 파일에서 텍스트를 추출한다. 테이블 구조도 보존."""
    pages = []
    page_count = 0
    try:
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    pages.append(text)

                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 5,
                    "join_tolerance": 5,
                })
                if not tables:
                    tables = page.extract_tables(table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                    })

                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(cell).strip() if cell else "" for cell in row]
                        if not any(cells):
                            continue
                        rows.append(" | ".join(cells))
                    if rows:
                        header = rows[0]
                        separator = " | ".join(["---"] * len(rows[0].split(" | ")))
                        body = "\n".join(rows[1:]) if len(rows) > 1 else ""
                        pages.append(f"[TABLE]\n{header}\n{separator}\n{body}\n[/TABLE]")

        full_text = "\n\n".join(pages)
        logger.info("PDF text extracted: %d pages, %d chars", page_count, len(full_text))
        return full_text
    except Exception as e:
        logger.error("PDF parsing failed: %s", e, exc_info=True)
        raise ValueError(f"PDF 파일을 읽을 수 없습니다: {e}")


def extract_text_from_docx(filepath):
    """DOCX 파일에서 텍스트를 추출한다."""
    from docx import Document
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 테이블도 추출
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                header = rows[0]
                separator = " | ".join(["---"] * len(rows[0].split(" | ")))
                body = "\n".join(rows[1:]) if len(rows) > 1 else ""
                paragraphs.append(f"[TABLE]\n{header}\n{separator}\n{body}\n[/TABLE]")

        full_text = "\n\n".join(paragraphs)
        logger.info("DOCX text extracted: %d chars", len(full_text))
        return full_text
    except Exception as e:
        logger.error("DOCX parsing failed: %s", e, exc_info=True)
        raise ValueError(f"DOCX 파일을 읽을 수 없습니다: {e}")
